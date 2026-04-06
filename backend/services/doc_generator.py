from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
from datetime import datetime, timedelta
import os


def _register_fonts():
    font_paths = [
        "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in font_paths:
        if os.path.exists(path):
            name = os.path.splitext(os.path.basename(path))[0]
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                return name
            except Exception:
                continue
    return "Helvetica"


def _weaving_label(weaving_type: str) -> str:
    return "трикотажные" if weaving_type == "трикотаж" else "швейные"


def _gender_label(target_group: str) -> str:
    if target_group == "adult_male":
        return "мужские"
    elif target_group == "adult_female":
        return "женские"
    return "детские"


def _main_material_label(compositions: list[dict]) -> str:
    if not compositions:
        return ""
    main = max(compositions, key=lambda c: c["percentage"])
    return main["material_name"]


def generate_docx(certificate: dict, client: dict, manufacturer: dict | None, products: list[dict]) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Header
    p = doc.add_paragraph("ЕВРАЗИЙСКИЙ ЭКОНОМИЧЕСКИЙ СОЮЗ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)

    doc_name = "СЕРТИФИКАТ СООТВЕТСТВИЯ" if certificate["doc_type"] == "CC" else "ДЕКЛАРАЦИЯ О СООТВЕТСТВИИ"
    h = doc.add_heading(doc_name, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Applicant
    ct = client.get("company_type", "ИП")
    cn = client.get("company_name", "")
    applicant_line = f"Заявитель: {ct} {cn}" if ct else f"Заявитель: {cn}"
    doc.add_paragraph(applicant_line)

    inn_line = f"ИНН {client.get('inn', '')}"
    if client.get("okpo"):
        inn_line += f" ОКПО {client['okpo']}"
    doc.add_paragraph(inn_line)

    if client.get("legal_address"):
        doc.add_paragraph(f"Место нахождения: Кыргызская Республика, {client['legal_address']}")

    contacts = ""
    if client.get("phone"):
        contacts += f"телефон: {client['phone']} "
    if client.get("email"):
        contacts += f"электронная почта: {client['email']}"
    if contacts:
        doc.add_paragraph(contacts)

    doc.add_paragraph(f"в лице {client.get('fio', '')}")

    # Product description
    if products:
        p0 = products[0]
        weaving = _weaving_label(p0.get("weaving_type", ""))
        gender = _gender_label(p0.get("target_group", ""))
        layer = p0.get("layer", 1)
        main_mat = _main_material_label(p0.get("compositions", []))
        tm_parts = [p.get("trademark", "") for p in products if p.get("trademark")]
        tm_str = ", ".join(f'«{t}»' for t in set(tm_parts) if t) if tm_parts else ""

        desc = f"заявляет, что Изделия {weaving} {layer} слоя {gender}"
        if main_mat:
            desc += f" из {main_mat}"
        if tm_str:
            desc += f", торговой марки {tm_str}"
        doc.add_paragraph(desc)

    if len(products) > 1:
        doc.add_paragraph(f"согласно приложению на {len(products) // 10 + 1} листе(ах)")

    doc.add_paragraph("Серийный выпуск")

    tnved_codes = list(set(p.get("tnved_code", "") for p in products if p.get("tnved_code")))
    if tnved_codes:
        doc.add_paragraph(f"Код ТН ВЭД: {', '.join(tnved_codes)}")

    # Manufacturer
    if manufacturer:
        mt = manufacturer.get("company_type", "ИП")
        mn = manufacturer.get("company_name", "")
        doc.add_paragraph(f"Изготовитель: {mt} {mn}")
        if manufacturer.get("production_address"):
            doc.add_paragraph(f"Адрес места осуществления деятельности: Кыргызская Республика, {manufacturer['production_address']}")
        elif manufacturer.get("legal_address"):
            doc.add_paragraph(f"Адрес: Кыргызская Республика, {manufacturer['legal_address']}")

    # TR TS
    tr = certificate["tr_ts"]
    if "007" in tr:
        tr_full = 'ТР ТС 007/2011 «О безопасности продукции, предназначенной для детей и подростков»'
    else:
        tr_full = 'ТР ТС 017/2011 «О безопасности продукции легкой промышленности»'
    doc.add_paragraph(f"Соответствует требованиям: {tr_full}")

    doc.add_paragraph("Декларация о соответствии принята на основании: Протокола испытаний")
    doc.add_paragraph("Схема декларирования: 3Д")
    doc.add_paragraph("Дополнительная информация: Изделия должны храниться в сухом, проветриваемом помещении в соответствии с правилами пожарной безопасности.")

    start = datetime.now()
    end = start + timedelta(days=365 * certificate["duration_years"])
    doc.add_paragraph(f"Декларация о соответствии действительна с {start.strftime('%d.%m.%Y')} по {end.strftime('%d.%m.%Y')} г. включительно.")

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Signature
    table = doc.add_table(rows=1, cols=4)
    cells = table.rows[0].cells
    cells[0].text = ""
    cells[1].text = "М.П."
    cells[2].text = "подпись"
    cells[3].text = f"{client.get('fio', '')}\n(Ф.И.О. заявителя)"

    # Appendix - page 2
    doc.add_page_break()
    p = doc.add_paragraph("ЕВРАЗИЙСКИЙ ЭКОНОМИЧЕСКИЙ СОЮЗ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2 = doc.add_heading("ПРИЛОЖЕНИЕ К ДЕКЛАРАЦИИ О СООТВЕТСТВИИ", level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Перечень продукции, на которую распространяется действие декларации о соответствии")

    # Products table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["№", "Код ТН ВЭД ТС", "Наименование и обозначение продукции", "Количество"]
    for i, h_text in enumerate(headers):
        table.rows[0].cells[i].text = h_text

    for idx, prod in enumerate(products, 1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = prod.get("tnved_code", "")
        comp_str = ", ".join(f"{c['material_name']} {c['percentage']}%" for c in prod.get("compositions", []))
        name = f"{prod.get('name', '')} ({comp_str})" if comp_str else prod.get("name", "")
        row.cells[2].text = name
        row.cells[3].text = ""

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf(certificate: dict, client: dict, manufacturer: dict | None, products: list[dict]) -> bytes:
    buf = BytesIO()
    font_name = _register_fonts()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    def text(t, x=2 * cm, size=11):
        nonlocal y
        c.setFont(font_name, size)
        c.drawString(x, y, t)
        y -= size + 6

    text("ЕВРАЗИЙСКИЙ ЭКОНОМИЧЕСКИЙ СОЮЗ", x=width / 2 - 4 * cm, size=10)
    y -= 5

    doc_name = "СЕРТИФИКАТ СООТВЕТСТВИЯ" if certificate["doc_type"] == "CC" else "ДЕКЛАРАЦИЯ О СООТВЕТСТВИИ"
    text(doc_name, x=width / 2 - 4 * cm, size=14)
    y -= 10

    ct = client.get("company_type", "ИП")
    cn = client.get("company_name", "")
    text(f"Заявитель: {ct} {cn}")
    text(f"ИНН {client.get('inn', '')}")
    if client.get("legal_address"):
        text(f"Адрес: {client['legal_address']}")
    text(f"в лице {client.get('fio', '')}")
    y -= 5

    if manufacturer:
        text(f"Изготовитель: {manufacturer.get('company_type', 'ИП')} {manufacturer.get('company_name', '')}")
        if manufacturer.get("production_address"):
            text(f"Адрес производства: {manufacturer['production_address']}")
    y -= 5

    text("ПРОДУКЦИЯ:", size=12)
    for prod in products:
        comps = ", ".join(f"{co['material_name']} {co['percentage']}%" for co in prod.get("compositions", []))
        text(f"  {prod.get('article', '')} — {prod.get('name', '')}")
        if comps:
            text(f"    Состав: {comps}")
        if prod.get("tnved_code"):
            text(f"    ТН ВЭД: {prod['tnved_code']}")
        if y < 3 * cm:
            c.showPage()
            y = height - 2 * cm

    y -= 10
    start = datetime.now()
    end = start + timedelta(days=365 * certificate["duration_years"])
    text(f"Действительна: {start.strftime('%d.%m.%Y')} — {end.strftime('%d.%m.%Y')}")
    text(f"Протоколов: {certificate['protocol_count']}")
    text(f"Стоимость: {certificate['total_price']:,.0f} сом")

    c.save()
    return buf.getvalue()
