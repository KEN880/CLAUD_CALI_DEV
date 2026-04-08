"""
Generate Макет ДС and Макет СС .docx files.
Templates match real Ares Consult documents exactly.
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ─── Helpers ───────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    """Set cell border. Usage: _set_cell_border(cell, top={"sz": 6, "val": "single"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tc.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = tc.makeelement(qn(f"w:{edge}"), {})
            tcBorders.append(element)
        for attr_name, attr_val in attrs.items():
            element.set(qn(f"w:{attr_name}"), str(attr_val))


def _no_borders(cell):
    """Remove all borders from cell."""
    for edge in ("top", "bottom", "start", "end"):
        _set_cell_border(cell, **{edge: {"sz": "0", "val": "none"}})


def _all_borders(cell, sz="4"):
    """Set all borders on cell."""
    border = {"sz": sz, "val": "single", "color": "000000"}
    _set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def _add_run(paragraph, text, font_size=12, bold=False, font_name="Times New Roman"):
    """Add a formatted run to paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    # Force font for Cyrillic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    return run


def _add_p(doc, alignment=None, space_after=None, space_before=None):
    """Add empty paragraph with formatting, return it for adding runs."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def _cell_run(cell, text, font_size=10, bold=False, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    _add_run(p, text, font_size=font_size, bold=bold)
    if alignment is not None:
        p.alignment = alignment
    return p


def _company_prefix(company_type: str) -> str:
    if company_type == "ИП":
        return "Индивидуальный предприниматель"
    return "Общество с ограниченной ответственностью"


def _weaving_adj(weaving: str) -> str:
    if weaving.lower() in ("трикотаж", "трикотажные"):
        return "трикотажные"
    return "швейные"


def _gender_adj(gender: str) -> str:
    return {"Мужской": "мужские", "Женский": "женские", "Детский": "детские"}.get(gender, "женские")


def _layer_word(layer: str) -> str:
    return {"1": "первого", "2": "второго", "3": "третьего"}.get(str(layer), f"{layer}")


def _format_trademarks(trademarks: str) -> str:
    parts = [t.strip() for t in trademarks.split(",") if t.strip()]
    if not parts:
        return ""
    return ", ".join(f"\u00ab{t}\u00bb" for t in parts)


def _signature_table(doc, name):
    """Add the standard signature table: [empty] [подпись] [М.П.] [ФИО]"""
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: '', '', 'М.П.', name
    _cell_run(table.rows[0].cells[2], "М.П.", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_run(table.rows[0].cells[3], name, font_size=12)

    # Row 1: '', 'подпись', '', '(Ф.И.О. заявителя)'
    _cell_run(table.rows[1].cells[1], "подпись", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_run(table.rows[1].cells[3], "(Ф.И.О. заявителя)", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # No borders on all cells
    for row in table.rows:
        for cell in row.cells:
            _no_borders(cell)

    # Widths
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(3)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(5)

    return table


def _registration_lines(doc):
    """Add registration number and date lines at bottom."""
    p = _add_p(doc, space_before=12)
    _add_run(p, "Регистрационный номер декларации о соответствии:", font_size=12)

    p = _add_p(doc, space_before=4)
    _add_run(p, "Дата регистрации декларации о соответствии: ", font_size=12)


# ─── Макет ДС (Declaration) ───────────────────────────

def generate_maket_ds(data: dict) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # Extract all data
    app_type = data.get("applicantType", "ИП")
    app_name = data.get("applicantName", "")
    app_inn = data.get("applicantInn", "")
    app_okpo = data.get("applicantOkpo", "")
    app_address = data.get("applicantAddress", "")
    app_phone = data.get("applicantPhone", "")
    app_email = data.get("applicantEmail", "")

    same = data.get("sameAsApplicant", False)
    if same:
        mfr_name = app_name
        mfr_address = app_address
        mfr_prod_address = app_address
    else:
        mfr_type = data.get("manufacturerType", "ИП")
        mfr_name = data.get("manufacturerName", "")
        mfr_address = data.get("manufacturerAddress", "")
        mfr_prod_address = data.get("manufacturerProductionAddress", "") or mfr_address

    weaving = data.get("productWeaving", "Трикотаж")
    layer = data.get("productLayer", "2")
    gender = data.get("productGender", "Женский")
    material = data.get("productMaterial", "")
    trademarks = data.get("productTrademarks", "")
    tr_ts = data.get("trTs", "017/2011")

    weaving_adj = _weaving_adj(weaving)
    gender_adj = _gender_adj(gender)
    layer_word = _layer_word(layer)
    tm_formatted = _format_trademarks(trademarks)
    app_prefix = _company_prefix(app_type)
    mfr_prefix = _company_prefix(mfr_type if not same else app_type)

    # ===== PAGE 1: DECLARATION BODY =====

    # Header
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_run(p, " ", font_size=10)
    _add_run(p, "ЕВРАЗИЙСКИЙ ЭКОНОМИЧЕСКИЙ СОЮЗ", font_size=13.5, bold=True)

    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_run(p, "ДЕКЛАРАЦИЯ О СООТВЕТСТВИИ", font_size=13.5, bold=True)

    # Заявитель
    p = _add_p(doc, space_after=2)
    _add_run(p, "Заявитель:", font_size=12, bold=True)
    _add_run(p, f" {app_prefix} {app_name}", font_size=12)

    # ИНН ОКПО
    p = _add_p(doc, space_after=2)
    _add_run(p, f"ИНН {app_inn} ОКПО {app_okpo} ", font_size=12)

    # Место нахождения
    p = _add_p(doc, space_after=2)
    _add_run(p, "Место нахождения:", font_size=12)
    _add_run(p, f" Кыргызская Республика, {app_address}", font_size=12)

    # Телефон, email
    p = _add_p(doc, space_after=2)
    _add_run(p, "телефон:", font_size=12)
    _add_run(p, f" {app_phone} ", font_size=12)
    _add_run(p, "электронная почта:", font_size=12)
    _add_run(p, f" {app_email}", font_size=12)

    # В лице
    p = _add_p(doc, space_after=4)
    _add_run(p, "в лице ", font_size=12, bold=True)
    _add_run(p, app_name, font_size=12)

    # Заявляет, что — product description
    material_text = f"из {material}" if material else ""
    tm_text = f", торговой марки {tm_formatted}" if tm_formatted else ""
    product_desc = f"Изделия {weaving_adj} {layer_word} слоя {gender_adj} {material_text}{tm_text}"

    p = _add_p(doc, space_after=2)
    _add_run(p, "заявляет, что ", font_size=12, bold=True)
    _add_run(p, product_desc.strip(), font_size=12)

    p = _add_p(doc, space_after=2)
    _add_run(p, "согласно приложению на 2 листе (ах) ", font_size=12)

    # Серийный выпуск
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    _add_run(p, "Серийный выпуск", font_size=12)

    # Код ТН ВЭД
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    _add_run(p, "Код ТН", font_size=12, bold=True)
    _add_run(p, " ВЭД", font_size=12, bold=True)
    _add_run(p, " ", font_size=12, bold=True)
    _add_run(p, "согласно приложению на 2 листе (ах)", font_size=12)

    # Изготовитель
    p = _add_p(doc, space_after=2)
    _add_run(p, "Изготовитель: ", font_size=12)
    _add_run(p, f"{mfr_prefix} ", font_size=12)
    _add_run(p, mfr_name, font_size=12)

    # Адрес
    p = _add_p(doc, space_after=4)
    _add_run(p, "Адрес места осуществления деятельности:", font_size=12)
    _add_run(p, f" Кыргызская Республика, {mfr_prod_address}", font_size=12)

    # Соответствует
    p = _add_p(doc, space_after=4)
    _add_run(p, "Соответствует требованиям:", font_size=12, bold=True)
    _add_run(p, f" ТР ТС {tr_ts} \u00abО безопасности продукции легкой промышленности\u00bb", font_size=12)

    # Принята на основании
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
    _add_run(p, "Декларация о соответствии принята на основании:", font_size=12, bold=True)
    _add_run(p, " Протокола испытаний", font_size=12)

    # Схема
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    _add_run(p, "Схема декларирования: 3Д", font_size=12)

    # Дополнительная информация
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)
    _add_run(p, "Дополнительная информация:", font_size=12, bold=True)
    _add_run(p, " Изделия должны храниться в сухом, проветриваемом помещении "
             "в соответствии с правилами пожарной безопасности в условиях, "
             "предотвращающих загрязнение, механические повреждения и действие "
             "солнечных лучей. ", font_size=12)

    # Действительна
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
    _add_run(p, "Декларация о соответствии действительна с даты регистрации по ", font_size=12)
    _add_run(p, "__", font_size=12)
    _add_run(p, " г. включительно.", font_size=12)

    # Signature table
    _signature_table(doc, app_name)

    # Registration lines
    _registration_lines(doc)

    # ===== PAGE 2: APPENDIX =====
    doc.add_page_break()

    # Header (with leading spaces to match template centering style)
    p = _add_p(doc, space_after=2)
    _add_run(p, "                                          ", font_size=12)
    _add_run(p, "ЕВРАЗИЙСКИЙ ЭКОНОМИЧЕСКИЙ СОЮЗ", font_size=12, bold=True)

    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_run(p, "     ПРИЛОЖЕНИЕ К ДЕКЛАРАЦИИ О СООТВЕТСТВИИ", font_size=12, bold=True)

    p = _add_p(doc, space_after=0)
    _add_run(p, "                         Перечень продукции, на которую   "
             "распространяется действие             ", font_size=12, bold=True)

    p = _add_p(doc, space_after=8)
    _add_run(p, "                                                         "
             "декларации о соответствии", font_size=12, bold=True)

    # Parse data
    tnved_lines = [l.strip() for l in data.get("tnvedCodes", "").strip().split("\n") if l.strip()]
    product_lines = [l.strip() for l in data.get("productList", "").strip().split("\n") if l.strip()]
    row_count = max(len(tnved_lines), len(product_lines), 3)

    # Product description header row (above numbered rows)
    product_header = f"Изделия {weaving_adj} {layer_word} слоя {gender_adj} {material_text}".strip()

    # Table with 4 columns: №, Код ТН ВЭД ТС, Наименование, Количество
    total_rows = 1 + 1 + row_count  # header + product description + data
    table = doc.add_table(rows=total_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["№", "Код ТН ВЭД ТС", "Наименование и обозначение продукции", "Количество\n(ед. измерения)"]
    for i, h in enumerate(headers):
        _cell_run(table.rows[0].cells[i], h, font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _all_borders(table.rows[0].cells[i], "6")

    # Product description row (row 1 — spans conceptually)
    _cell_run(table.rows[1].cells[0], "", font_size=10)
    _cell_run(table.rows[1].cells[1], "", font_size=10)
    _cell_run(table.rows[1].cells[2], product_header, font_size=10)
    _cell_run(table.rows[1].cells[3], "", font_size=10)
    for cell in table.rows[1].cells:
        _all_borders(cell)

    # Data rows
    for row_idx in range(row_count):
        row = table.rows[row_idx + 2]
        code = tnved_lines[row_idx] if row_idx < len(tnved_lines) else ""
        product = product_lines[row_idx] if row_idx < len(product_lines) else ""

        _cell_run(row.cells[0], str(row_idx + 1), font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_run(row.cells[1], code, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_run(row.cells[2], product, font_size=10)
        _cell_run(row.cells[3], "", font_size=10)
        for cell in row.cells:
            _all_borders(cell)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(9)
        row.cells[3].width = Cm(3.3)

    _add_p(doc, space_after=8)

    # Signature table on appendix too
    _signature_table(doc, app_name)

    # Registration lines
    _registration_lines(doc)

    # Serialize
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─── Макет СС (Certificate) ──────────────────────────

def generate_maket_cc(data: dict) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # Extract data
    app_type = data.get("applicantType", "ИП")
    app_name = data.get("applicantName", "")
    app_inn = data.get("applicantInn", "")
    app_address = data.get("applicantAddress", "")
    app_phone = data.get("applicantPhone", "")
    app_email = data.get("applicantEmail", "")
    app_prefix = _company_prefix(app_type)

    same = data.get("sameAsApplicant", False)
    if same:
        mfr_name = app_name
        mfr_address = app_address
    else:
        mfr_name = data.get("manufacturerName", "")
        mfr_address = data.get("manufacturerProductionAddress", "") or data.get("manufacturerAddress", "")

    weaving = data.get("productWeaving", "Трикотаж")
    layer = data.get("productLayer", "2")
    gender = data.get("productGender", "Женский")
    trademarks = data.get("productTrademarks", "")
    tr_ts = data.get("trTs", "017/2011")

    weaving_adj = _weaving_adj(weaving)
    gender_adj = _gender_adj(gender)
    layer_word = _layer_word(layer)
    tm_formatted = _format_trademarks(trademarks) if trademarks else "\u00ab\u00bb"

    # ===== PAGE 1: CERTIFICATE BODY =====

    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    _add_run(p, "№ ЕАЭС KG …..", font_size=11)

    p = _add_p(doc, space_after=8)
    _add_run(p, "ОРГАН ПО СЕРТИФИКАЦИИ …..", font_size=11, bold=True)

    # ЗАЯВИТЕЛЬ — no OKPO in CC
    p = _add_p(doc, space_after=6)
    _add_run(p, f"ЗАЯВИТЕЛЬ  {app_prefix} {app_name}, ИНН {app_inn}; "
             f"Место нахождения: Кыргызская Республика, {app_address} "
             f"Место осуществления деятельности: "
             f"телефон: {app_phone} электронная почта: {app_email}", font_size=11)

    # ИЗГОТОВИТЕЛЬ
    p = _add_p(doc, space_after=6)
    _add_run(p, f"ИЗГОТОВИТЕЛЬ  {mfr_name} "
             f"Место осуществления деятельности: {mfr_address}", font_size=11)

    # ПРОДУКЦИЯ
    p = _add_p(doc, space_after=6)
    _add_run(p, f"ПРОДУКЦИЯ    Изделия {weaving_adj} {layer_word} слоя {gender_adj}, "
             f"торговой марки {tm_formatted}, "
             f"согласно приложению, на 2 листе(ах), серийный выпуск;", font_size=11)

    # КОД ТН ВЭД
    p = _add_p(doc, space_after=6)
    _add_run(p, "КОД ТН ВЭД ЕАЭС  согласно приложению на 2 листе(ах)", font_size=11)

    # СООТВЕТСТВУЕТ
    p = _add_p(doc, space_after=6)
    _add_run(p, f"СООТВЕТСТВУЕТ ТРЕБОВАНИЯМ  ТР ТС {tr_ts} "
             f"\u00abО безопасности продукции легкой промышленности\u00bb.", font_size=11)

    # Expert fills
    p = _add_p(doc, space_after=6)
    _add_run(p, "СЕРТИФИКАТ ВЫДАН НА ОСНОВАНИИ   ……..", font_size=11)
    p = _add_p(doc, space_after=6)
    _add_run(p, "ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ  …………", font_size=11)
    p = _add_p(doc, space_after=12)
    _add_run(p, "СРОК ДЕЙСТВИЯ C …… ПО …….", font_size=11)

    # Signature blocks
    p = _add_p(doc, space_after=2)
    _add_run(p, "Руководитель (уполномоченное лицо) органа по сертификации"
             "    _____________     М.П.  __________________", font_size=10)
    p = _add_p(doc, space_after=8)
    _add_run(p, "                                  (подпись)          (Ф.И.О.)", font_size=9)
    p = _add_p(doc, space_after=2)
    _add_run(p, "Эксперт (эксперт-аудитор)"
             " (эксперты (эксперты-аудиторы)) _____________     М.П.     _________________", font_size=10)
    p = _add_p(doc, space_after=4)
    _add_run(p, "                                  (подпись)             (Ф.И.О.)", font_size=9)

    # ===== PAGE 2: APPENDIX =====
    doc.add_page_break()

    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_run(p, "к сертификату соответствия № ЕАЭС KG ….", font_size=11, bold=True)
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _add_run(p, "Перечень конкретной продукции,", font_size=11, bold=True)
    p = _add_p(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _add_run(p, "на которую распространяется действие сертификата соответствия", font_size=11, bold=True)

    # Product description
    p = _add_p(doc, space_after=6)
    _add_run(p, f"Изделия {weaving_adj} {layer_word} слоя, торговой марки {tm_formatted}", font_size=11)

    # Table: №, Код ТН ВЭД ЕАЭС, Наименование
    tnved_lines = [l.strip() for l in data.get("tnvedCodes", "").strip().split("\n") if l.strip()]
    product_lines = [l.strip() for l in data.get("productList", "").strip().split("\n") if l.strip()]
    row_count = max(len(tnved_lines), len(product_lines), 5)

    table = doc.add_table(rows=row_count + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["№", "Код ТН ВЭД\nЕАЭС", "Наименование и обозначение продукции, ее изготовитель"]
    for i, h in enumerate(headers):
        _cell_run(table.rows[0].cells[i], h, font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _all_borders(table.rows[0].cells[i], "6")

    for row_idx in range(row_count):
        row = table.rows[row_idx + 1]
        code = tnved_lines[row_idx] if row_idx < len(tnved_lines) else ""
        product = product_lines[row_idx] if row_idx < len(product_lines) else ""

        _cell_run(row.cells[0], str(row_idx + 1), font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_run(row.cells[1], code, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_run(row.cells[2], product, font_size=10)
        for cell in row.cells:
            _all_borders(cell)

    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(12)

    # Signature blocks
    _add_p(doc, space_after=12)
    p = _add_p(doc, space_after=2)
    _add_run(p, "Руководитель (уполномоченное лицо) органа по сертификации"
             "    _____________     М.П.  __________________", font_size=10)
    p = _add_p(doc, space_after=8)
    _add_run(p, "                                  (подпись)          (Ф.И.О.)", font_size=9)
    p = _add_p(doc, space_after=2)
    _add_run(p, "Эксперт (эксперт-аудитор)"
             " (эксперты (эксперты-аудиторы)) _____________     М.П.     _________________", font_size=10)
    p = _add_p(doc, space_after=0)
    _add_run(p, "                                  (подпись)             (Ф.И.О.)", font_size=9)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
